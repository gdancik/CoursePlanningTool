from flask import Flask, render_template, request, send_file, session, jsonify
from docx import Document
from flask_login import LoginManager, UserMixin, login_user, login_required, logout_user, current_user
from flask_session import Session

from gs_editor import gsEditor

from flask_cors import CORS

import pandas as pd

import requests
import io
import random

import course_planning as cp

''' Create app and login manager '''
app = Flask(__name__)
app.secret_key = "my secret key"

CORS(app, resources = {r"/api/*": 
    {"origins": [
        "http://127.0.0.1:3000",
        "http://localhost:3000",
        "need to add domain here"
    ]}},
    supports_credentials = True
)

login_manager = LoginManager()
login_manager.init_app(app)

app.config["SESSION_PERMANENT"] = True     # Sessions expire when the browser is closed
app.config["SESSION_TYPE"] = "filesystem"     # Store session data in files

Session(app)

def get_gs_editor() :
    #return gsEditor('annie')
    #return gsEditor(session['user'])
    return gsEditor(current_user.id)

''' Classes and functions to handle login '''
class User(UserMixin):
    def __init__(self, id, name):
      self.id = id
      self.name = name

@login_manager.user_loader
def load_user(user_id):
    # need to look up user based on user_id
    return User(user_id, user_id)

@app.route('/api/test_login/', methods = ['GET'])
def test_login() :
    error = 'Error: url must be in format /?test_login/user=user&password=password', 400
    if (len(request.args) != 2 or
        'user' not in request.args or 
        'password' not in request.args): 
            return error
    
    username = request.args['user']
    password = request.args['password']
    
    if password != 'password' : 
        return 'invalid password', 401
        
    user = User(username, username)
    login_user(user)
   
    gs = get_gs_editor()
    gs.create_sheet()

    return jsonify(user = username)    

@app.route('/api/logout/')
def logout():
    logout_user()
    session.clear()
    return 'User is logged out, now you cannot access <a href = "/api/hello/">/api/hello/</a>' 

@app.route('/profile/')
def profile():
   if current_user.is_authenticated :
      return f'<p>ID: {current_user.id}</p><p>Name: {current_user.name}</p>'
   return '<p>You are currently not logged in. To log in, go to <a href = "/login/">/login/</a>'

# route to the homepage
@app.route('/')
def index() :
    s = '''
    <h1> Course Planning Tool Homepage</h1>
    <ul>
    <li> <a href = '/test_login/?user=annie&password=password'>Test Login</a> </li>
    <li> <a href = '/logout/'>Logout</a> </li>
    <li> <a href = '/api/hello/'>Hello</a> </li>
    <li> <a href = '/profile/'>Profile</a> </li>
    <li> <a href = '/valid_inputs/'>Valid Inputs </a> </li>
    </ul>
    '''
    return s

''' 
The 'session' object allows you to store information specific to a user. We
may be able to accomplish the same thing using flask_login and current_user
''' 

@app.route('/api/get_session_number/')
def get_session_number():
    session['number'] = random.randint(0,100)
    return 'A random number has been assigned; go to /show_session_number/ to view'

@app.route('/api/show_session_number/')
def test():
    number = session.get('number', None) 
    if number :
      return f'The number is : {number}'

    return 'To get a number, go to /get_session_number/'


''' 
API calls should return a jsonified message.
The 'hello' endpoint requires a login, while the 'hi' one does not
'''

@app.route('/api/hello/')
@login_required
def hello():
    return jsonify(message="Hello there from Flask!")
    #if current_user :
    #  return {'message': current_user.id}
    #return 'no current user!'

@app.route('/api/hi/')
def hi():
    return jsonify(message="Hi there from Flask!")

@app.route('/api/valid_inputs/')
def valid_inputs():
    return jsonify(cp.columns)

''' Generate a word document! '''
@app.route('/generate/', methods=['GET', 'POST'])
def generate():

    if request.method == 'POST':
        name = request.form.get('name')
        message = request.form.get('message')

        # Create Word document in memory
        doc = Document()
        doc.add_heading(f"Message from {name}", 0)
        doc.add_paragraph(message)

        # Save to a BytesIO stream
        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)

        return send_file(
            file_stream,
            as_attachment=True,
            download_name=f"{name}_message.docx",
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )

    return render_template('form.html')


''' Preview syllabus '''
@app.route('/api/preview/', methods=['POST'])
def preview():

    try:
        data = request.get_json()
        name = data.get('name')
        message = data.get('message')

    except Exception as e:
        return jsonify({"error": str(e)}), 400

    #return jsonify({'name': name, 'message': message})

    # Create Word document in memory
    doc = Document()
    doc.add_heading(f"Message from {name}", 0)
    doc.add_paragraph(message)

    # Save to a BytesIO stream
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)

    return send_file(
        file_stream,
        as_attachment=True,
        download_name=f"message_from_{name}.docx",
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )


'''Calls the getValue function from the gs_editor.py module'''

@app.route('/api/getValue/', methods=['POST'])
@login_required
def getValue():
    gs = get_gs_editor()
    try:
        data = request.get_json()
        course_id = data.get('course_id')
        columns = data.get('list_of_columns')
        sheet_name = data.get('sheet_name')
        if not course_id or not columns:
            return jsonify({"error": "Missing one or more required fields"}), 400

        # Call the getValue function
        gs.set_sheet_name(sheet_name)
        sheet = gs.getValue(course_id, columns,)

        # Check if sheet is None
        if sheet is None:
            return jsonify({"error": "No data returned from getValue"}), 500

        # Return the result as JSON
        return jsonify(sheet)

    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/updateValue/', methods=['POST'])
@login_required
def updateValue():
    gs = get_gs_editor()
    try:
        
        data = request.get_json()
        course_id = data.get('course_id')
        columns = data.get('list_of_columns')
        sheet_name = data.get('sheet_name')

        if not course_id or not columns:
            return jsonify({"error": "Missing one or more required fields"}), 400

        # Call the getValue function
        gs.set_sheet_name(sheet_name)
        gs.updateValue(course_id, columns)
        return jsonify('Function called successfully')
    except Exception as e:
        return jsonify({"error": str(e)}), 500

# TO DO: login should be required
@app.route('/api/getCourses/', methods=['POST'])
def getCourses():
        data = request.get_json()
        user_id = data.get('user')
        
        if not user_id:
            return jsonify({"error": "Missing one or more required fields"}), 400


        df = pd.DataFrame({'a': [1,2,3],
                           'b': [4,9,20]})
        
        return jsonify(df.to_dict(orient = 'records'))

#TO DO: login should be required
@app.route('/api/getNewCourseId/', methods=['POST'])
def getNewCourseId():
        data = request.get_json()
        user_id = data.get('user')
        
        if not user_id:
            return jsonify({"error": "Missing one or more required fields"}), 400

        return jsonify(course_id = '4')


'''Calls the getValue function from the gs_editor.py module'''
@app.route('/api/getSheet/', methods=['POST'])
@login_required
def getSheet():
    gs = get_gs_editor()
    try:
        print('requesting data')
        data = request.get_json()
        print(data) 
        sheet_name = data.get('sheet_name')
        if not sheet_name:
            return jsonify({"error": "sheet_name is required"}), 400

        # Call the getValue function
        gs.set_sheet_name(sheet_name)
        sheet = gs.read_sheet()

        # Check if sheet is None
        if sheet is None:
            return jsonify({"error": "No data returned from getSheet"}), 500

        # Return the result as JSON
        return jsonify(sheet.to_dict(orient = 'records'))

    except Exception as e:
        return jsonify({"error": str(e)}), 500

'''Calls the create_sheet function to share the current sheet'''
@app.route('/api/shareSheet/', methods=['POST'])
def shareSheet():
    gs = get_gs_editor()
    try:
        data = request.get_json()                
        email = data.get('email')
        if not email:
            return jsonify({"error": "email is required"}), 400

        # Call the create_sheet method        
        id = gs.create_sheet(email = email)
        url = 'https://docs.google.com/spreadsheets/d/' + id
        # Return the result as JSON
        return jsonify({'id': id, 'url': url})

    except Exception as e:
        return jsonify({"error": str(e)}), 500

'''Admin page'''
@app.route('/admin/')
def admin():

    return '<h2>Under construction</h2>'

    gs = get_gs_editor()
    url = f'https://docs.google.com/spreadsheets/d/{gs.id}'
    
    page = f'''
    <h2> Admin Page </h2>
    <ul>
    <li>Sheet name: {gs.sheet_name} </li>
    <li>Sheet id: {gs.id} </li>
    <li>url: <a href = "{url}">{url}</a></li>
    <li>API count: {gs.api_count}
    </ul>
    '''
    return page

if __name__ == '__main__':
    app.run(debug=True)

