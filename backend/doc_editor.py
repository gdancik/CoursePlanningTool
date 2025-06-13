from docx import Document
import pandas as pd
import tabulate
from python_docx_replace import docx_replace, docx_blocks

#TODO Need to change to be more algorithmically efficient
#Currently it is O(n^2) because it iterates through the paragraphs for each key in the dictionary

def replaceTextInParagraph(doc, fr_dict):
    '''
    Replaces text in a Word document based on a dictionary of replacements.
    Args:
        doc: Word document object.
        fr_dict (dict): Dictionary containing text to be replaced as keys and their replacements as values.
        block_ls (list, optional): List of blocks to be removed. Defaults to None.
    '''

    docx_replace(doc, **fr_dict)
        
def removeBlocks(doc, block_ls):
    '''
    Removes specified blocks from a Word document.
    Args:
        doc: Word document object.
        block_ls (list): List of blocks to be removed.
    Returns:
        None
    '''
    
    for i in block_ls:
        options = {i: False}
    docx_blocks(doc, **options)
 
def printParagraphs(doc):
    '''
    Prints all paragraphs in a Word document.
    Args:
        doc: Word document object.
    Returns:
        None
    '''
   
    # Iterate through the paragraphs in the document and print them
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:  # Only print non-empty paragraphs
            print(text)

def getParagraph(doc, i):
    '''
    Retrieves a specific paragraph from a Word document.
    Args:
        doc: Word document object.
        i (int): Index of the paragraph to retrieve.
    Returns:
        str: Text of the specified paragraph.
    '''
    
    paragraph = doc.paragraphs[i].text.strip()
    return paragraph
    
def print_table(t):
    '''
    Prints a DataFrame in a tabular format.
    Args:
        t (pd.DataFrame): DataFrame to be printed.
    Returns:
        None
    '''
    #Check if t is a DataFrame
    if not isinstance(t, pd.DataFrame):
        raise ValueError("Input must be a pandas DataFrame.")

    # Convert the DataFrame to a tabular format using tabulate
    table = tabulate.tabulate(t, headers='keys', tablefmt='grid')
    # Print the table
    print(table)

def printTables(doc):
    '''
    Prints all tables in a Word document in a tabular format.
    Args:
        doc: Word document object.
    Returns:
        None
    '''
    # Open the document
    
    # Iterate through the tables in the document and print them
    for i in range(len(doc.tables)):
        # Convert the table to a DataFrame
        table = getTable(doc, i)
        # Print the DataFrame in a tabular format
        print_table(table)

def getTable(doc, i):
    '''
    Retrieves a specific table from a Word document and converts it to a DataFrame.
    Args:
        doc: Word document object.
        i (int): Index of the table to retrieve.
    Returns:
        pd.DataFrame: DataFrame containing the table data.
    '''
    # Open the document

    table = doc.tables[i]
   
    # Initialize an empty list to store rows
    rows = []
    # Iterate over each row in the table
    for row in table.rows:
        # Initialize an empty list to store cell contents for the current row
        row_data = []

        # Iterate over each cell in the current row
        for cell in row.cells:
            # Strip the text content of the cell and add it to the row_data list
            row_data.append(cell.text.strip())

        # Add the row_data list to the rows list
        rows.append(row_data)

    # Create a DataFrame from the list of rows
    df = pd.DataFrame(rows)
    return df


