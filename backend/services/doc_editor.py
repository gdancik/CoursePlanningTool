from docx import Document
import pandas as pd
import tabulate
from python_docx_replace import docx_replace, docx_blocks
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy
import re
import docx
import logging


#TODO Need to change to be more algorithmically efficient
#Currently it is O(n^2) because it iterates through the paragraphs for each key in the dictionary

def replaceTextInParagraph(doc, fr_dict):
    '''
    Replaces text in a Word document based on a dictionary of replacements.
    Args:
        doc: Word document object.
        fr_dict (dict): Dictionary containing text to be replaced as keys and their replacements as values.
        block_ls (list, optional): List of blocks to be removed. Defaults to None.
    '''
    logging.info('Replacing text in document')
    docx_replace(doc, **fr_dict)


''' Convert run to dictionary '''
def run_to_dict(r) :
    return {    
            'text': r.text,
            'style': r.style,
            'bold': r.bold,
            'italic': r.italic
            }

'''
create_runs
    - run: a run possibly containing markdown
    - type: should be one of 'bold', 'italic', or 'both'

    Returns a paragraph with formatted runs, or None if formatting not done
'''
def create_runs(run, type):
    
    pattern = r'\*\*(.*?)\*\*'
    if type == 'both' :
        pattern = r'\*\*\*(.*?)\*\*\*'
    elif type == 'italic' :
        pattern = r'\*(.*?)\*'
    
    text = run.text    
    parts = re.split(pattern, text)
    
    n = len(parts)
    if n == 1 :
        return None
    
    new_doc = Document()
    new_p = new_doc.add_paragraph()
    
    run.text = ""     # Clear original run
    for i in range(n) :        
        run1 = new_p.add_run(parts[i])
        run1.style = run.style
        if i % 2 == 1 :
            if type == 'bold' or type == 'both' :
                run1.bold = True
            if type == 'italic' or type == 'both' :
                run1.italic = True
    return new_doc

''' Extract and return a list of all runs from a document'''
def extract_runs (doc) :
    rlist = []
    for p in doc.paragraphs :
        for r in p.runs :
            rlist.append(r)
    return rlist

''' 
    Reformats paragraphs containing markdown for type = 'bold', 'italic', or 'both'
    Needs to be called in this order: 'both', 'bold', 'italic' to handle repeated *'s
'''

def reformat_paragraphs(doc) :
    reformat_paragraphs_by_type(doc, 'both')
    reformat_paragraphs_by_type(doc, 'bold')
    reformat_paragraphs_by_type(doc, 'italic')

def reformat_paragraphs_by_type(doc, type):
    for p in doc.paragraphs :
        runList = []
        formatted = False
        for r in p.runs :
            # need to keep list of all runs in case we need to reformat
            runList.append(r)
        
            # break into new runs if necessary       
            new_doc = create_runs(r, type)  
            if new_doc :
                runList += extract_runs(new_doc)           
                formatted = True     

        # if reformatting, re-create the paragraph
        if formatted :
            p.clear()
            for r in runList :
                rd = run_to_dict(r)
                nr = p.add_run(rd.get('text'))
                nr.style = rd.get('style')
                nr.bold = rd.get('bold')
                nr.italic = rd.get('italic')


    
def add_text_in_runs(paragraph, run_positions, start, end, bold=False):
    """Add text between start and end indices using original run styles."""
    remaining_start = start
    while remaining_start < end:
        # Find the run that contains this start
        for run_start, run_end, orig_run in run_positions:
            if run_start <= remaining_start < run_end:
                run_relative_start = remaining_start - run_start
                run_relative_end = min(end - run_start, run_end - run_start)
                text_segment = orig_run.text[run_relative_start:run_relative_end]
                new_run = paragraph.add_run(text_segment)
                new_run.style = orig_run.style
                if bold:
                    new_run.bold = True
                remaining_start += len(text_segment)
                break


def replaceFormattedTextInParagraph(doc):
    '''
    Replaces markdown-formatted text in a Word document with formatted text (bold, italic).
    Args:
        doc: Word document object.
    '''    

    reformat_paragraphs(doc)
      
        
def removeBlocks(doc, block_ls):
    '''
    Removes specified blocks from a Word document.
    Args:
        doc: Word document object.
        block_ls (list): List of blocks to be removed.
    Returns:
        None
    '''
    logging.info('Removing blocks from doc')
    options = {}
    for i in block_ls:
        options[i] = False
    docx_blocks(doc, **options)

def removeBlockTags(doc, block_ls):
    '''
    Removes specified block tags from a given document.

    Args:
        doc (Document): The document object from which block tags will be removed.
        block_ls (list of str): A list of block tag identifiers that need to be removed from the document.

    Returns:
        None: This function does not return a value but modifies the document in place.
    '''
    options = {}
    for i in block_ls:
        options[i] =True
    docx_blocks(doc, **options)
 
def printParagraphs(doc):
    '''
    Prints all paragraphs in a Word document.
    Args:
        doc: Word document object.
    Returns:
        None
    '''
    logging.info('Printing paragraphs')
    # Iterate through the paragraphs in the document and print them
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:  # Only print non-empty paragraphs
            print(text)

def getParagraph(doc, i):
    '''
    Retrieves a specific paragraph from a Word document.
    Args:
        doc: Word document object.
        i (int): Index of the paragraph to retrieve.
    Returns:
        str: Text of the specified paragraph.
    '''
    logging.info(f'Fetching paragraph {i}')
    paragraph = doc.paragraphs[i].text.strip()
    return paragraph
    
def print_table(t):
    '''
    Prints a DataFrame in a tabular format.
    Args:
        t (pd.DataFrame): DataFrame to be printed.
    Returns:
        None
    '''
    logging.info('Printing table')
    #Check if t is a DataFrame
    if not isinstance(t, pd.DataFrame):
        raise ValueError("Input must be a pandas DataFrame.")

    # Convert the DataFrame to a tabular format using tabulate
    table = tabulate.tabulate(t, headers='keys', tablefmt='grid')
    # Print the table
    print(table)

def printTables(doc):
    '''
    Prints all tables in a Word document in a tabular format.
    Args:
        doc: Word document object.
    Returns:
        None
    '''
    logging.info('Printing all tables in the doc')
    
    # Iterate through the tables in the document and print them
    for i in range(len(doc.tables)):
        # Convert the table to a DataFrame
        table = getTable(doc, i)
        # Print the DataFrame in a tabular format
        print_table(table)

def getTable(doc, i):
    '''
    Retrieves a specific table from a Word document and converts it to a DataFrame.
    Args:
        doc: Word document object.
        i (int): Index of the table to retrieve.
    Returns:
        pd.DataFrame: DataFrame containing the table data.
    '''
    logging.info(f'Fetching table {i}')

    table = doc.tables[i]
   
    # Initialize an empty list to store rows
    rows = []
    # Iterate over each row in the table
    for row in table.rows:
        # Initialize an empty list to store cell contents for the current row
        row_data = []

        # Iterate over each cell in the current row
        for cell in row.cells:
            # Strip the text content of the cell and add it to the row_data list
            row_data.append(cell.text.strip())

        # Add the row_data list to the rows list
        rows.append(row_data)

    # Create a DataFrame from the list of rows
    df = pd.DataFrame(rows)
    return df

# GD: this does not work with hyperlinks, use 
# copy_paragraph_with_html_before instead 
# for syllabus policies
def copy_paragraph_before(source_paragraph, target_paragraph):
    '''
    Copies a paragraph from one document and inserts it before a specified paragraph in another document,
    preserving text, formatting, and hyperlinks.
    Args:
        source_paragraph: Paragraph object from the source document to be copied.
        target_paragraph: Paragraph object in the target document before which the new paragraph will be inserted.
    
    '''
    def add_hyperlink(paragraph, text, url):
        def get_or_create_hyperlink_style(d):
            """If this document had no hyperlinks so far, the builtin
            Hyperlink style will likely be missing and we need to add it.
            There's no predefined value, different Word versions
            define it differently.
            This version is how Word 2019 defines it in the
            default theme, excluding a theme reference.
            """
            if "Hyperlink" not in d.styles:
                if "Default Character Font" not in d.styles:
                    ds = d.styles.add_style("Default Character Font",
                                            docx.enum.style.WD_STYLE_TYPE.CHARACTER,
                                            True)
                    ds.element.set(docx.oxml.shared.qn('w:default'), "1")
                    ds.priority = 1
                    ds.hidden = True
                    ds.unhide_when_used = True
                    del ds
                hs = d.styles.add_style("Hyperlink",
                                        docx.enum.style.WD_STYLE_TYPE.CHARACTER,
                                        True)
                hs.base_style = d.styles["Default Character Font"]
                hs.unhide_when_used = True
                hs.font.color.rgb = docx.shared.RGBColor(0x05, 0x63, 0xC1)
                hs.font.underline = True
                del hs

            return "Hyperlink"
        # This gets access to the document.xml.rels file and gets a new relation id value
        part = paragraph.part
        r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

        # Create the w:hyperlink tag and add needed values
        hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
        hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

        # Create a new run object (a wrapper over a 'w:r' element)
        new_run = docx.text.run.Run(
            docx.oxml.shared.OxmlElement('w:r'), paragraph)
        new_run.text = text

        # Set the run's style to the builtin hyperlink style, defining it if necessary
        new_run.style = get_or_create_hyperlink_style(part.document)
        # Alternatively, set the run's formatting explicitly
        # new_run.font.color.rgb = docx.shared.RGBColor(0, 0, 255)
        # new_run.font.underline = True

        # Join all the xml elements together
        hyperlink.append(new_run._element)
        paragraph._p.append(hyperlink)
        return hyperlink
    
    # Extract text from the source paragraph
    text = source_paragraph.text
    link_index = None

    # Insert a new paragraph before the target paragraph in the target document
    # This new paragraph will initially contain the text from the source paragraph
    new_paragraph = target_paragraph.insert_paragraph_before(text)

    # Copy the style of the source paragraph to the new paragraph
    if source_paragraph.style:
        new_paragraph.style = source_paragraph.style

    # Clear the target paragraph as it will be rebuilt with the new content
    target_paragraph.clear()

    # Copy formatting from the source paragraph to the new paragraph
    if source_paragraph.runs:
        # Clear existing runs in the new paragraph to start fresh
        new_paragraph.clear()

        # Check if there is a hyperlink in the source paragraph and get its index and text
        for child in source_paragraph._element.iterchildren():
            # logging.debug(f'Paragraph child: {child}')#debug
            # Iterate through the XML elements of the source paragraph to find hyperlinks
            if child.tag.endswith('hyperlink'):
                # Log that a hyperlink has been found for debugging purposes
                logging.debug('Found link')#debug
                # Store the index of the hyperlink
                link_index = source_paragraph._element.index(child)
                # Extract the text of the hyperlink
                link_text = source_paragraph._element.getchildren()[link_index].text
                logging.debug(f'link text: {link_text}, link run: {link_index}')#debug

        # Copy each run from the source paragraph to the new paragraph
        for i, run in enumerate(source_paragraph.runs):
            # Add a new run to the new paragraph with the same text as the source run
            new_run = new_paragraph.add_run(run.text)

            # Copy the formatting properties (bold, italic, underline) from the source run
            new_run.bold = run.bold
            new_run.italic = run.italic
            new_run.underline = run.underline
            new_run.font.color.rgb = run.font.color.rgb
            new_run.font.size = run.font.size
            new_run.font.name = run.font.name
            # # If the current run index matches the hyperlink index, append the hyperlink text
            if i + 1 == link_index:
                if link_text.startswith("https://"):
                    add_hyperlink(new_paragraph,link_text,link_text)
                else:
                    link_text_url = f'https://{link_text}'
                    add_hyperlink(new_paragraph,link_text,link_text_url)
                   
                
            # if i + 1 == link_index:
            #     new_run.text += link_text

            # Check if the current run is part of a hyperlink using lxml for debug
            current_run = run._element
            logging.debug(f'current_run: {current_run}, run: {i}')  # Debugging log
                    
                
def copy_paragraph_with_html_before(source_doc, target_doc, target_para):
    '''Copies all text in source_doc to before target_para
    Must handle hyperlinks by copying relationships and remapping rIds
    '''
    
    # Hyperlinks use rId references — we need to port them over so the links resolve
    rel_id_map = {}
    for rel in source_doc.part.rels.values():
        if "hyperlink" in rel.reltype:        
        # Add the relationship to the target doc and capture the new rId
            new_rId = target_doc.part.relate_to(rel.target_ref, rel.reltype, is_external=True)
            rel_id_map[rel.rId] = new_rId

    for para in reversed(source_doc.paragraphs):
        para_element = copy.deepcopy(para._element)

        # Remap any rId references in hyperlink elements to the new doc's rIds
        for hyperlink in para_element.findall(f'.//{qn("w:hyperlink")}'):
            old_rId = hyperlink.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id')
            if old_rId and old_rId in rel_id_map:
                hyperlink.set(
                '{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id',
                rel_id_map[old_rId]
            )

        target_para._element.addnext(para_element)
