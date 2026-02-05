import requests
import json
from bs4 import BeautifulSoup
from docx import Document
from htmldocx import HtmlToDocx
from docx.shared import RGBColor, Pt, Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn 

import backend.services.doc_editor as de
import backend.services.firestore_editor as fse
import backend.services.course_planning as cp
import re
import logging


#%%
def get_webpage(url: str) :
    ''' 
    Returns the text from url, if valid.
    Args:
        url (str): The URL of the webpage to scrape.
    Returns:
        BeautifulSoup object: Parsed HTML content of the webpage.
    '''
    logging.info(f'Fetching webpage')
    try :        
        r = requests.get(url)
    except Exception as error:
        # handle the exception
        logging.error("An exception occurred:", error)

    if r.status_code != 200 :
        raise(f'Error with connection to {url}\n, status code is {r.status_code}')

    return BeautifulSoup(r.text, 'html.parser')

def extractFromAccordian(soup):
    '''
    Extracts headers and content from accordions in the given BeautifulSoup object.
    Args:
        soup (BeautifulSoup): Parsed HTML content of the webpage.
    Returns:
        dict: A dictionary where keys are headers and values are the corresponding content.
    '''
    logging.info(f'Extracting text from accordians')
    headers_and_content = {}
    classes = soup.find_all(class_='accordion')
    for item in classes:
        rows = item.find_all(class_='row')
        for i in rows:
            header = i.find('h3')
            content = i.find_all('div', class_='content')
            headers_and_content[header] = content 
    return headers_and_content

def getStatements(url,selected_statements:list = None):
    '''
    Gets the syllabus statements from the given URL. by searching for the accordion classes then finding the rows and extracting the header and content.
    Args:
        url (str): The URL of the webpage to scrape.
        selected_statements (list, optional): A list of strings to filter the statements by their headers. If None, all statements are returned.
    Returns:
        dict: A dictionary where keys are headers and values are lists of content for the corresponding statements.
    '''

    soup = get_webpage(url)

    logging.info(f'Getting selected syllabus statements')

    #get all statements
    all_statements = extractFromAccordian(soup)

    #Filter only the statements selected (Will have to search header for string)
    statments = {} #all_statements

    logging.debug(f'selected_statements: {selected_statements},type:{type(selected_statements)}')#debug
    if selected_statements:
        statments = {}
        for i in selected_statements:
            # print(f'Current statment: {i}')
            for key, val in all_statements.items():
                if i.lower() in str(key).lower():
                    # print('Current key: {key}')
                    statments[key] = val
    #return selected statements
    return statments

def create_syllabus_statment_page(doc,url: str,selected_statements=None):
    '''
    Creates a syllabus statement page in the given Word document by fetching statements from a the syllabus statement website and adding them to the document.
    Args:
        doc (Document): The Word document object to which the syllabus statements will be added.
        url (str): The URL of the webpage to scrape for syllabus statements.
        selected_statements (list, optional): A list of strings to filter the statements by their headers. If None, all statements are returned.
    Returns:
        None
    '''
    logging.info(f'Creating statement page')

    logging.debug(f'selected_statements: {selected_statements}, type:{type(selected_statements)}')#debug

   

    x = getStatements(url,selected_statements)

    for header, content in x.items():
        
        # get header and remove whitespace after any closing tag
        header_string = str(header).strip()
        header_string = re.sub(r'(</[^>]+>)\s+', r'\1', header_string)
        content_string = str(list(content)[0])
        html_to_word_htmldocx(doc,header_string)
        html_to_word_htmldocx(doc,content_string)

def html_to_word_htmldocx(doc,html_content: str):
    '''
    Converts HTML content to a Word document using the HtmlToDocx module and adds it to the specified document.
    Args:
        doc (Document): The Word document object to which the HTML content will be added.
        html_content (str): The HTML content to be converted and added to the document.
    Returns:
        None
    '''

    logging.debug(f'Formatting html to doc')
    cleaned_html = _strip_whitespace_from_div(html_content)
    new_parser = HtmlToDocx()
    new_parser.add_html_to_document(cleaned_html,doc)

def _strip_whitespace_from_div(html_content: str):
    '''
    Strips unnecessary whitespace from HTML content, specifically around <p> and <div class="content"> tags.
    Args:
        html_content (str): The HTML content to be cleaned.
    Returns:
        str: Cleaned HTML content with unnecessary whitespace removed.
    '''
    stripped_html = re.sub(r'\s+(<p>)', r'\1', html_content)
    stripped_html = re.sub(r'(<\/p>)\s+', r'\1', stripped_html)  
    stripped_html = re.sub(r'(<\/div class="content">)\s+', r'\1', stripped_html)
    return stripped_html

def add_table_to_doc(doc, table_list:list, merge: list = None, header = True):
    '''
    Adds a table to the given Word document using the provided list of lists.
    Args:
        doc (Document): The Word document object to which the table will be added.
        table_list (list of lists): A list of lists representing the table data.
        merge (list of list): A list of list where each sublist contains three elements:
            - The row number (1-indexed) where the merge should occur.
            - The first cell to merge (1-indexed).
            - The second cell to merge (1-indexed).
        header (bool, optional): If True, the first row of table_list is treated as the header row. Defaults to True.
    Returns:
        table (Table): The created table object.
    '''

    logging.info(f'Adding table to document')
    table = doc.add_table(rows=len(table_list), cols=len(table_list[0]),style = "Table Grid")
   
    if merge:
        logging.debug('Merging Cells')
        #iterate through list
        for i in merge:
            #for each list the first val is the row num and the other 2 are the cells to combine
            row = table.rows[i[0]-1]
            cell1= row.cells[i[1]-1]
            cell2= row.cells[i[2]-1]
            cell1.merge(cell2)

    for i, row in enumerate(table_list):
        for j, cell_data in enumerate(row):
            table.cell(i, j).text = str(cell_data)

    if header == True:
        style_table_header(table)

    # change formatting for Grade and Assessment tables
    if table_list[0][0] == 'Grade' or table_list[0][0] == 'Assessment':
        table.autofit = False
        for column in table.columns:
            for cell in column.cells:
                cell.width = Inches(2)
   
    return table
    
def center_table_text(table):
    '''
    Centers the text in all cells of a specified table in the given Word document.
    Args:
        doc (Document): The Word document object containing the table.
        table (Table): The table object whose cell text will be centered.
    Returns:
        None
    '''
    logging.debug(f'Centering table text')
    for row in table.rows:
        for cell in row.cells:
            cell.paragraphs[0].alignment = 1  # 1 corresponds to center alignment

def style_table_borders(table):
    '''
    Styles the borders of a specified table in the given Word document.
    Args:
        doc (Document): The Word document object containing the table.
        table (Table): The table object whose borders will be styled.
    Returns:
        None
    '''
    logging.debug(f'Styling table borders')
    for row in table.rows:
            for cell in row.cells:
                tcPr = cell._element.tcPr  # Get the table cell properties element
                if tcPr is None:
                    tcPr = OxmlElement('w:tcPr')
                    cell._element.append(tcPr)

            
                tcBorders = tcPr.find(qn('w:tcBorders'))
                if tcBorders is None:
                    tcBorders = OxmlElement('w:tcBorders')
                    tcPr.append(tcBorders)

                #Couldn't get the loop for this working so I just ran it for every cell - Sencere
                top_border = tcBorders.find(qn('w:top'))
                if top_border is None:
                    top_border = OxmlElement('w:top')
                    tcBorders.append(top_border)
                top_border.set(qn('w:val'), 'nil')

                #only set first column's left border to nil
                if cell._element.getparent().index(cell._element) == 0:
                    # Set the left border style to 'nil' (no border)
                    left_border = tcBorders.find(qn('w:left'))
                    if left_border is None:
                        left_border = OxmlElement('w:left')
                        tcBorders.append(left_border)
                    left_border.set(qn('w:val'), 'nil')

                #only set last column's right border to nil
                if cell._element.getparent().index(cell._element) == len(row.cells) - 1:
                    # Set the right border style to 'nil' (no border)
                    right_border = tcBorders.find(qn('w:right'))
                    if right_border is None:
                        right_border = OxmlElement('w:right')
                        tcBorders.append(right_border)
                    right_border.set(qn('w:val'), 'nil')

def style_table_text(table):
    '''
    Changes the font, size, and color of all text in a specified table in the given Word document.
    Args:
        table (Table): The table object whose font will be changed.
    Returns:
        None
    '''
    logging.debug(f'Styling table text')
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
                run.font.size = Pt(11)
                run.font.name = 'Calibri'
                run.font.color.rgb = RGBColor(0, 0, 0)
                
def style_table_header(table):
    '''
    Changes the font of the header row in a specified table in the given Word document.
    Args:
        table (Table): The table object whose header font will be changed.
    Returns:
        None
    '''
    logging.debug(f'styling table heaer')
    header_row = table.rows[0]
    for cell in header_row.cells:
        for paragraph in cell.paragraphs:
            run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
            run.bold = True
            run.font.color.rgb = RGBColor(32, 44, 92)
            run.font.size = Pt(12)
            # paragraph.alignment = 1  # Center align the header text

def add_styled_table(doc, table_list:list, header = True, center = False):
    '''
    Adds a styled table to the given Word document using the provided list of lists.
    Args:
        doc (Document): The Word document object to which the table will be added.  
        table_list (list of lists): A list of lists representing the table data.
        header (bool, optional): If True, the first row of table_list is treated as the header row. Defaults to True.
        center (bool, optional): If True, the text in the table will be centered. Defaults to False.
    Returns:
        table (Table): The created and styled table object.
    '''
    # Make table
    table = add_table_to_doc(doc, table_list, False)
    # Style table
    style_table_borders(table)
    if center == True:
        center_table_text(table)
    style_table_text(table)
    if header == True:
        style_table_header(table)
    return table

def generate_grading_policies(doc, policies:list):
    '''
    Generates a section in the Word document for grading policies, formatting each policy with a title and description.
    Args:
        doc (Document): The Word document object to which the grading policies will be added.
        policies (list of dict): A list of dictionaries where each dictionary contains a policy title and its description.
    Returns:
        None
    '''
    for i in policies:
        #If its an outcome
        if len(i) == 2:           
            #doc.add_paragraph(f'{i['title']}: {i['description']}')
            #return
            items = list(i.items())            
            key_label, key, value = items[0][0], items[0][1], items[1][1]
            #title
            title = key

            key_label, title = "", i['title']   # GD

            paragraph = doc.add_paragraph(f'{title.strip()}')
            run_title = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
            run_title.font.color.rgb = RGBColor(32, 44, 92)
            run_title.font.size = Pt(12)
            run_title.font.name = 'Calibri'
            run_title.bold = True
            run_title.italic = True
            #description
            description = value
            description = i['description']      # GD
            paragraph = doc.add_paragraph(description)
            run_description = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
            run_description.font.color.rgb = RGBColor(0, 0, 0)
            run_description.font.size = Pt(11)
            run_description.font.name = 'Calibri'
            run_description.bold = False
            run_description.italic = False

        if len(i) == 3:
            items = list(i.items())
            key_label, key = items[0][0], items[0][1]
            points_label,points =  items[1][0],items[1][1]
            value =  items[2][1]
            #title
            title = key

            key_label, title = "", i['title']  # GD
            points_label, points = "", i['rightValue'] # GD
            
            paragraph = doc.add_paragraph(f'{title} ({points})')
            run_title = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
            run_title.font.color.rgb = RGBColor(32, 44, 92)
            run_title.font.size = Pt(12)
            run_title.font.name = 'Calibri'
            run_title.bold = True
            run_title.italic = True

            #description
            description = value
            description = i['description']      # GD
          
            paragraph = doc.add_paragraph(description)
            run_description = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
            run_description.font.color.rgb = RGBColor(0, 0, 0)
            run_description.font.size = Pt(11)
            run_description.font.name = 'Calibri'
            run_description.bold = False
            run_description.italic = False
           
                
def generate_syllabus(doc: object, course_id:str, sheet_name: str, syllabus_statment_webpage_url:str='https://www.easternct.edu/center-for-teaching-learning-and-assessment/syllabus-statements/index.html'):
    """ 
    Generates a syllabus document by replacing placeholders with actual values.

    Args:
        doc (Document): The Word document object to be modified.
        course_id (str): The ID of the course for which the syllabus is generated.
        sheet_name (str): The name of the document to retrieve course data from.
        syllabus_statment_webpage_url (str, optional): The URL of the syllabus statement webpage. Defaults to a specific URL.
    Returns:
        Title(str): The title of the syllabus document, which is a combination of course information.
    """

    def table_placeholder_replacement(doc, paragraph, placeholder_text, table_list):
        if placeholder_text in paragraph.text:
            table = add_styled_table(doc, table_list)
            paragraph.clear()
            paragraph._p.addnext(table._tbl)
            return True
        return False
    
    # Retrieve course data using fsEditor
    fs = fse.fsEditor(sheet_name)
    column_names = cp.columns

    # do not use getValue, since these returns None values
    # fr_dict = fs.getValue(course_id, column_names)
    #fr_dict = fs.getValue(course_id, column_names)
    fr_dict = fs.getCourse(course_id)
    
    # convert from list to string for specified columns
    for i in cp.to_string:
        if i in fr_dict:
            #if fr_dict[i]:
            #    print(f'convert list to string: fr_dict["{i}"] = {fr_dict[i]}')
                #fr_dict[i] = json.loads(fr_dict[i])         
            fr_dict[i] = ''.join(fr_dict[i]) if isinstance(fr_dict[i], list) else fr_dict[i]

     # Process dictionary to remove "_checkboxes" suffix from keys
    for key, value in list(fr_dict.items()):
        if key.endswith('_checkboxes'):
            new_key = key[:-11]
            fr_dict[new_key] = value
            # Optionally, remove the old key if you want to replace it
            del fr_dict[key]
    logging.debug('Removing "_checkboxes" from the column names')

    # Process dictionary to remove "_syllabus" suffix from keys
    syllabus_col = {key[:-9]: value for key, value in fr_dict.items() if key.endswith('_syllabus')}
    logging.debug('Removing "_syllabus" from the column names')
    


    # Handle items that are displayed as numbered lists
    # TO DO: Use appropriate naming convention so this is automatic

    placeholder = "${course_materials}"
    item_string = fr_dict.get('course_materials_syllabus', None)
    items = [item.strip() for item in item_string.split('\n') if item.strip()] if item_string else None
    generate_numbered_list(doc, placeholder, items, p_style = 'CPT Numbered List')

    placeholder = "${ lo_syllabus_json }"
    item_json = fr_dict.get('lo_syllabus_json', None) 

    items = None
    titles = None
    if item_json :
        items = []        
        titles = []
        for i, item in enumerate(item_json) :
            s = f'Learning outcome {i+1} (LO{i+1}'
            if item['title'] :
                s += ', ' + item['title'] + ')'
            titles.append(s)
            items.append(': ' + item['description'])
                
    generate_numbered_list(doc, placeholder, items, titles, p_style = 'Bullet List')

    ## Add Assignment Table and Descriptions
    placeholder = '${assmt_assignments_syllabus_json}'
    item_json = fr_dict.get('assmt_assignments_syllabus_json', None) 

    if item_json: 
        titles = [d.get('title', '') for d in item_json]
        descriptions = [d.get('description', '') for d in item_json]
        percentages = [d.get('rightValue','') for d in item_json]
        
        # must be first since placeholder is included twice in the syllabus  
        assessment_table = [['Assessment', 'Points or Percentage']] + list(zip(titles, percentages))

        for p in doc.paragraphs :
            if placeholder in p.text :
                table_placeholder_replacement(doc, p, placeholder, assessment_table)
                break

        titles = [t  + ': ' for t in titles]
        generate_numbered_list(doc, placeholder, descriptions, titles)
    
    
    # Find and replace are handled here
    de.replaceTextInParagraph(doc, syllabus_col)

    # Process dictionary to remove "_json" suffix from keys
    json_columns = {key[:-5]: value for key, value in fr_dict.items() if key.endswith('_json')}
    logging.debug('Removing "_json" from the column names')
    logging.debug(f'json_columns:{json_columns}')
    

    # Process dictionary to handle table placeholders
    tables_col = {key[:-5]: value for key, value in fr_dict.items() if key.endswith('_list')}
    logging.debug('Removing "_list" from the column names')


    #print('fr_dict = ', fr_dict)

    #policy placeholder handling
    policies = fr_dict.get('policy', [])
    policies += fr_dict.get('resources', ['Accommodations for Students with Disabilities'])

    '''
    try:
        logging.debug(f'string before conversion to literal:{policies}')
        if policies == None:
            pass    
        else:
            policies = json.loads(policies)
    except ValueError as e:
        logging.error(f"Error converting string to literal: {e}")
    '''
        

    logging.debug(f'policies: {policies}, type:{type(policies)}')#debug

    syllabus_statment_page = Document()
    create_syllabus_statment_page(syllabus_statment_page,syllabus_statment_webpage_url,policies) 


    # Iterate through paragraphs and replace placeholders
    for paragraph in doc.paragraphs:
        #Add policies
        if '<policy_statements></policy_statements>'in paragraph.text: 
            logging.debug('Policy Placeholder found') #debug
            for source_paragraph in syllabus_statment_page.paragraphs:
                de.copy_paragraph_before(source_paragraph,paragraph)
        

        for key, value in json_columns.items():
            if key in paragraph.text:
                logging.debug(f'key: {key}') #debug
                logging.debug(f'Value: {value}') #debug
                if value == None:
                    logging.debug('Value is None, skipping') #debug
                    break
                list_of_dicts = json_columns.get(key)
                logging.debug(f'List of dicts: {list_of_dicts}') #debug
                try:
                    logging.debug(f'skip -- string before conversion to literal list_of_dicts:{list_of_dicts}')
                    # GD: this was causing an erro
                    #list_of_dicts= json.loads(list_of_dicts)
                except ValueError as e:
                    logging.error(f"Error converting string to literal: {e}")
                logging.debug(f'list_of_dicts: {list_of_dicts}, type:{type(list_of_dicts)}')#debug

                generate_grading_policies_page = Document()
                generate_grading_policies(generate_grading_policies_page,list_of_dicts)
                for source_paragraph in generate_grading_policies_page.paragraphs:
                    de.copy_paragraph_before(source_paragraph,paragraph)
        
        
        # Go through table placeholders
        logging.debug('Processing table placeholders')
        for key, value in tables_col.items():
            logging.debug(f'placeholder name: {key}')
            if value:
                try:
                    # Clean and evaluate the string to convert it to a list
                    cleaned_value = value.rstrip('\'')
                    try:
                        evaluated_value = json.loads(cleaned_value)
                    except ValueError as e:
                        logging.error(f"Error converting string to literal: {e}")
                    logging.debug(f'Table_value: {evaluated_value}, type:{type(evaluated_value)}')#debug
                    if table_placeholder_replacement(doc, paragraph, key, evaluated_value):
                        logging.info('Table placeholder replaced')
                except (ValueError, SyntaxError) as e:
                    logging.error(f"Error evaluating string for key {key}: {e}")


    # Blocks are removed
    removeTags = []
    removeBlocks = []

    if fr_dict.get('times2_syllabus') == "":
        removeBlocks.append('time2')
    else:
        removeTags.append('time2')
    
    if fr_dict.get('instructor_title_syllabus') == "":
        removeBlocks.append('title')
    else:
        removeTags.append('title')

    if fr_dict.get('instructor_pronouns_syllabus') == "":
        removeBlocks.append('pronouns')
    else:
        removeTags.append('pronouns')
    
    
    if fr_dict.get('form_of_address_syllabus') == "":
        removeBlocks.append('foa')
    else:
        removeTags.append('foa')

    if fr_dict.get('phone_syllabus') == "":
        removeBlocks.append('phone')
    else: 
        removeTags.append('phone')
    
    if fr_dict.get('instructor_additional_information') == "":
        removeBlocks.append('instructor_additional_information')
    else:
        removeTags.append('instructor_additional_information')

    if fr_dict.get('hip'):
        removeTags.append('hip')        
    else: 
        removeBlocks.append('HIP')
    
    # always remove -- can we do this for all of them???
    removeBlocks.append('policy_statements')
    
    de.removeBlocks(doc,removeBlocks)
    de.removeBlockTags(doc,removeTags)
    #Replace any formatted text that is in markdown
    
    de.replaceFormattedTextInParagraph(doc)
    
    title = ""
    title += str(syllabus_col.get('subj_code', None))
    title += str(syllabus_col.get('crse_number',None))
    title += "_" + str(syllabus_col.get('term',None))

    return title

'''Replaces placeholder with a numbered list in doc'''
def generate_numbered_list(doc, placeholder, items, titles = None, p_style = None):
    if items :        
   
        for i, paragraph in enumerate(doc.paragraphs):
            if placeholder in paragraph.text:
                # Remove the placeholder paragraph
                #p = paragraph._element
                anchor = paragraph                               

                # Insert numbered list items (reverse order preserves order)
                for item_idx, item in reversed(list(enumerate(items))):
                    new_p = doc.paragraphs[i].insert_paragraph_before()

                    if p_style :
                        new_p.style = p_style
                                        

                    if titles :
                        #r = new_p.add_run(titles[item_idx])
                        r = new_p.add_run(titles[item_idx])
                        r.style = 'List Item Title'
                                        
                    new_p.add_run(item)
                                                                       
                anchor._element.getparent().remove(anchor._element)
                break
          
       
    
    
