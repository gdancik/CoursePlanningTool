from flask import request, jsonify, send_file, render_template
from flask_login import login_required, login_user
from docx import Document
from datetime import datetime
import pandas as pd
import io

from backend.models.login import User
import backend.services.course_planning as cp
from backend.services.app_services import get_gs_editor

from . import test_api_bp

@test_api_bp.route('/')
def api_test_home() :
    return 'test_api'

@test_api_bp.route('test_login/', methods = ['GET'])
def test_login() :
    error = 'Error: url must be in format /?test_login/user=user&password=password', 400
    if (len(request.args) != 2 or
        'user' not in request.args or 
        'password' not in request.args): 
            return error
    
    username = request.args['user']
    password = request.args['password']
    
    if password != 'password' : 
        return 'invalid password', 401
        
    user = User(username, username)
    login_user(user)

    gs = get_gs_editor()
    gs.create_sheet()

    return jsonify(user = username)    

@test_api_bp.route('/test_data/', methods = ['GET'])
@login_required
def test_data() :
    '''
    Generates data for course_id = 'test'
    Data has format column_name + "1"
    '''

    try :
        gs = get_gs_editor()
        gs.create_sheet()

        course_id = 'test'

        d = {col: col + '1' for col in cp.columns if col != 'course_id'}

        gs.updateValue(course_id, d)
    
    except Exception as e :
        return jsonify({"error": str(e)}), 400
        
    return jsonify(course_id = course_id, status = 'success'), 200    


''' 
API calls should return a jsonified message.
The 'hello' endpoint requires a login, while the 'hi' one does not
'''
@test_api_bp.route('hello/')
@login_required
def hello():
    return jsonify(message="Hello there from Flask!")
    #if current_user :
    #  return {'message': current_user.id}
    #return 'no current user!'

@test_api_bp.route('hi/')
def hi():
    return jsonify(message="Hi there from Flask!")

@test_api_bp.route('valid_inputs/')
def valid_inputs():
    return jsonify(cp.columns)

''' Generate a word document! '''
@test_api_bp.route('generate/', methods=['POST'])
def generate():

    if request.method == 'POST':
        name = request.form.get('name')
        message = request.form.get('message')

        # Create Word document in memory
        doc = Document()
        doc.add_heading(f"Message from {name}", 0)
        doc.add_paragraph(message)

        # Save to a BytesIO stream
        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)

        return send_file(
            file_stream,
            as_attachment=True,
            download_name=f"{name}_message.docx",
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )

    return render_template('form.html')


'''Calls the create_sheet function to share the current sheet'''
@test_api_bp.route('shareSheet/', methods=['POST'])
@login_required
def shareSheet():
    gs = get_gs_editor()
    try:
        data = request.get_json()                
        email = data.get('email')
        if not email:
            return jsonify({"error": "email is required"}), 400

        # Call the create_sheet method        
        id = gs.create_sheet(email = email)
        url = 'https://docs.google.com/spreadsheets/d/' + id
        # Return the result as JSON
        return jsonify({'id': id, 'url': url})

    except Exception as e:
        return jsonify({"error": str(e)}), 500


@test_api_bp.route('error/', methods = ['GET'])
def test_error() :
    error = 'Error: url must be in format /error/?code=code&error=error', 400
    if (len(request.args) != 2 or
        'code' not in request.args or 
        'error' not in request.args): 
            return error
    
    code = int(request.args['code'])
    error = request.args['error']
    
    return jsonify(error = error), code    
