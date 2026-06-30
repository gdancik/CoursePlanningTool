from docx import Document

# Change this import to your actual file/module name
from backend.services.syllabus_generator import add_styled_table


def main():
    normal_table = [
        ["Date", "Day", "Unit and Theme/Topic", "Learning Outcomes Addressed", "Reading/Assignments Due"],
        ["01/21/2026", "Wednesday", "Normal text with spaces should look unchanged.", "", ""],
    ]

    long_word_table = [
        ["Date", "Day", "Unit and Theme/Topic", "Learning Outcomes Addressed", "Reading/Assignments Due"],
        [
            "01/22/2026",
            "Thursday",
            "wassdwadwadwadadawdawswdwassdwadwadwadadawdawswdwassdwadwadwadadawdawswd",
            "",
            "",
        ],
    ]

    doc = Document()

    doc.add_heading("Normal table", level=1)
    add_styled_table(doc, normal_table)

    doc.add_paragraph("")

    doc.add_heading("Long unbroken word table", level=1)
    add_styled_table(doc, long_word_table)

    doc.save("test_break_long_words.docx")

    print("Saved test_break_long_words.docx")


if __name__ == "__main__":
    main()