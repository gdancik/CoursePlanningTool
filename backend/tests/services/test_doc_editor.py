import sys
import os

# Add the greatgrandparent directory to the Python path -- required for direct testing of this file
sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), '../../../')))

import pytest
import backend.services.doc_editor as de
from docx import Document
import pandas as pd

WIP = 0
if WIP == 1:
    document = 'TestFiles/(Test)Doc1.docx'
else: 
    document = 'backend/tests/services/TestFiles/(Test)Doc1.docx'


doc= Document(document)
def test_replaceTextInParagraph():
    test_dict ={
        'Name': 'Fetty',
        'Age': '30',
        'Hobby': 'cooking'
                }
    de.replaceTextInParagraph(doc,test_dict)
    modified_doc = doc
    para = ""
    for paragraph in modified_doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            para += text + "\n"
    assert para == 'Hello my name is Fetty, I am 30 years old.\nMy favorite hobby is cooking.\n<test>Test block</test>\n'
    
def test_printParagraphs(capsys):
    doc= Document(document)
    de.printParagraphs(doc)
    captured = capsys.readouterr()
    assert captured.out == 'Hello my name is ${Name}, I am ${Age} years old.\nMy favorite hobby is ${Hobby}.\n<test>Test block</test>\n'

def test_getParagraph():
    doc= Document(document)
    para = de.getParagraph(doc, 0)
    para = para.strip()
    assert para == 'Hello my name is ${Name}, I am ${Age} years old.'

def test_printTables(capsys):
    de.printTables(doc)
    captured = capsys.readouterr()
    x = '''
+----+-----+-------+------+
|    | 0   | 1     | 2    |
+====+=====+=======+======+
|  0 | Hi  | this  | is   |
+----+-----+-------+------+
|  1 | my  | table | that |
+----+-----+-------+------+
|  2 | I   | test  | with |
+----+-----+-------+------+
        '''
    assert captured.out.strip() == x.strip()

def test_getTable():
    x = [
        ['Hi', 'this', 'is'],
        ['my', 'table', 'that'],
        ['I', 'test', 'with']
    ]

    expected_df = pd.DataFrame(x)
    y = de.getTable(doc,0)
    pd.testing.assert_frame_equal(y, expected_df)
    
def test_removeBlocks():
    doc= Document(document)
    de.removeBlocks(doc,['test'])
    para = ""
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            para += text + "\n"
    assert para == 'Hello my name is ${Name}, I am ${Age} years old.\nMy favorite hobby is ${Hobby}.\n'
    pass

def test_replaceFormattedTextInParagraph():
    if WIP == 1:
        markdown_test_doc = Document('TestFiles/(Test)Markdown.docx')
    else:
        markdown_test_doc = Document('backend/tests/services/TestFiles/(Test)Markdown.docx')

    de.replaceFormattedTextInParagraph(markdown_test_doc)
    modified_doc = markdown_test_doc
    para = ""
    for paragraph in modified_doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            para += text + "\n"
    print(para)
    assert para == 'This is bold\nthis is italics\nThis is both\n'

def test_copy_paragraph_before():
    if WIP == 1:
        hyperlink_test_doc = Document('TestFiles/(Test)Hyperlink.docx')
    else:
        hyperlink_test_doc = Document('backend/tests/services/TestFiles/(Test)Hyperlink.docx')
    test_doc = Document()
    test_doc.add_paragraph()
    para = hyperlink_test_doc.paragraphs[0]
    de.copy_paragraph_before(para,test_doc.paragraphs[0])
    assert test_doc.paragraphs[0].text == para.text
