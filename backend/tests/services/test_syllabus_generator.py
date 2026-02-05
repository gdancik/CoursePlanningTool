import sys
import os

# Add the greatgrandparent directory to the Python path -- required for direct testing of this file
sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), '../../../')))

import pytest
import backend.services.syllabus_generator as sf
import backend.services.firestore_stats as fs_stats

from docx import Document

syllabus_webpage = 'https://www.easternct.edu/center-for-teaching-learning-and-assessment/syllabus-statements/index.html'
WIP = 0

@pytest.fixture
def db_setup():
    fs_stats.create_tables()
    yield

def test_get_webpage():
    url = "http://example.com" 
    soup = sf.get_webpage(url)

    assert soup.title is not None
    assert "Example Domain" in soup.title.string

def test_getStatements():
   
    statements_dict = sf.getStatements(syllabus_webpage,['Diversity Statement']) 

    statement_keys = list(statements_dict.keys())
    statement_values = list(statements_dict.values())
    expected_diversity_statement = '''Eastern Connecticut State University values the diversity of its students, faculty, and staff. Differences in race, ethnicity, national origin, class, religion, learning styles, gender, gender identity and expression, sexual orientation, age, ideology, and other aspects of human variation and characterization, including but not limited to those protected by law and CSCU policies, enrich the educational experiences and social and intellectual development of students and create a rich cultural environment. Eastern is committed to ensuring that regardless of their differences, all members of the Eastern community are challenged to achieve their full potential and are supported in their pursuit of that goal in a campus environment that is free from discrimination and harassment.'''
    
    assert 'Diversity Statement' in str(statement_keys[0])
    assert expected_diversity_statement.strip() in str(statement_values[0])

def test_create_syllabus_statment_page(db_setup):
    doc= Document()
    selected_statements= ['Academic Success Center']
    
    sf.create_syllabus_statment_page(doc,syllabus_webpage,selected_statements)
    expected_text = 'The Academic Success Center (ASC) is located on the ground floor of the library and assists Eastern students in realizing their highest level of achievement possible. We encourage all students to take advantage of the many resources offered through the following ASC offices: Advising Center, Global Learning, Opportunity Programs, Career Services, Writing Center, Math Achievement Center, and Tutoring & Learning Strategies.' 
    paragraphs = []
    for paragraph in doc.paragraphs:
        paragraphs.append(paragraph.text)
    assert 'Academic Success Center (ASC)' in paragraphs[0]
    assert expected_text in paragraphs[1]

#Testing that the code doesnt crash
def test_generate_syllabus(db_setup):
    if WIP == 1:
        doc= Document('TestFiles/(Test)Policy.docx')
    else:
        doc= Document('backend/tests/services/TestFiles/(Test)Doc1.docx')
    
    title = sf.generate_syllabus(doc,'0fBVjCYVsxV73bA00a8R','annie')
    # doc.save('test_syllabus.docx')
    assert title != None

def test_add_table_to_doc():
    doc= Document()
    data =  [
        ['Header1', 'Header2'],
        ['Row1Col1', 'Row1Col2'],
        ['Row2Col1', 'Row2Col2']
    ]
    sf.add_table_to_doc(doc, data)

    paragraphs = [para.text for para in doc.paragraphs]
    table_found = False
    for table in doc.tables:
        if len(table.rows) == 3 and len(table.columns) == 2:
            table_found = True
            break
    assert table_found

def test_style_table_borders():
    
    doc= Document()
    data =  [
        ['Header1', 'Header2'],
        ['Row1Col1', 'Row1Col2']
    ]
    sf.add_table_to_doc(doc, data)
    table = doc.tables[0]
    sf.style_table_borders(table)

def test_generate_grading_policies():
    if WIP == 1:
        doc = Document('TestFiles/(Test)Policy.docx')
    else:
        doc = Document('backend/tests/services/TestFiles/(Test)Policy.docx')

    data = [
        {'title':'40-49','description':'Failing'},
        {'title':'50-59','description':'Passing'},
    ]
    sf.generate_grading_policies(doc, data)
    paragraphs = [para.text for para in doc.paragraphs]
    assert 'Passing' in paragraphs[len(paragraphs)-1]
   
