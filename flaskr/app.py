from flask import Flask, render_template, request, send_file, session, jsonify
from docx import Document
import io
import random

app = Flask(__name__)
app.secret_key = "super secret key"

@app.route('/', methods=['GET', 'POST'])
def index():

    # generate a random number each time this page is reloaded
    # the number is stored in session and will be accessible across pages
    session['name'] = random.randint(0,1000)

    if request.method == 'POST':
        name = request.form.get('name')
        message = request.form.get('message')

        # Create Word document in memory
        doc = Document()
        doc.add_heading(f"Message from {name}", 0)
        doc.add_paragraph(message)

        # Save to a BytesIO stream
        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)

        return send_file(
            file_stream,
            as_attachment=True,
            download_name=f"{name}_message.docx",
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )

    return render_template('form.html')


@app.route('/test/')
def test():
    name = session.get('name', 'unassigned')
    return render_template('test.html', name = name)

@app.route('/api/hello')
def hello():
    return jsonify(message="Hello from Flask!")

if __name__ == '__main__':
    app.run(debug=True)

